#!/usr/bin/env python3
"""
Routes Module - API Endpoints
Không init db ở đây, chỉ import và sử dụng
"""

from flask import Blueprint, request, jsonify, send_file, current_app
from werkzeug.utils import secure_filename
from models import db, Document, Attachment, ChatMessage
from ai_service import AIService
from datetime import datetime
import os
import PyPDF2
from docx import Document as DocxDocument
import shutil
from pathlib import Path
import logging

# Setup logging
logger = logging.getLogger(__name__)

# ✅ Tạo Blueprint - KHÔNG init db ở đây
api_bp = Blueprint('api', __name__, url_prefix='/api')


# ============ FILE EXTRACTION ============

def extract_file_content(file_path, filename):
    """Trích xuất nội dung từ file"""
    try:
        ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''

        if ext == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()

        elif ext == 'pdf':
            content = []
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        content.append(text)
            return '\n'.join(content) if content else '[PDF không có nội dung]'

        elif ext in ['doc', 'docx']:
            try:
                doc = DocxDocument(file_path)
                content = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        content.append(para.text)
                return '\n'.join(content) if content else '[Document không có nội dung]'
            except:
                return '[Không thể đọc file Word]'

        return '[Loại file không được hỗ trợ]'

    except Exception as e:
        logger.error(f"Error extracting content: {str(e)}")
        return f'[Lỗi trích xuất: {str(e)}]'


# ============ API ENDPOINTS ============

@api_bp.route('/health', methods=['GET'])
def health_check():
    """Kiểm tra sức khỏe API"""
    try:
        # Test database connection
        db.session.execute('SELECT 1')

        return jsonify({
            'success': True,
            'status': 'OK',
            'message': 'Server is running',
            'timestamp': datetime.utcnow().isoformat()
        }), 200
    except Exception as e:
        return jsonify({
            'success': False,
            'status': 'ERROR',
            'error': str(e)
        }), 500


@api_bp.route('/documents', methods=['GET'])
def get_documents():
    """Lấy danh sách tất cả văn bản"""
    try:
        docs = Document.query.order_by(Document.created_at.desc()).all()
        return jsonify({
            'success': True,
            'documents': [doc.to_dict() for doc in docs]
        }), 200
    except Exception as e:
        logger.error(f"Error getting documents: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@api_bp.route('/documents/<doc_id>', methods=['GET'])
def get_document(doc_id):
    """Lấy chi tiết một văn bản"""
    try:
        doc = Document.query.get(doc_id)
        if not doc:
            return jsonify({'success': False, 'error': 'Document not found'}), 404

        return jsonify({
            'success': True,
            'document': {
                'id': doc.id,
                'title': doc.title,
                'content': doc.content,
                'document_type': doc.document_type,
                'document_number': doc.document_number,
                'sender': doc.sender,
                'receiver': doc.receiver,
                'date_received': doc.date_received.isoformat() if doc.date_received else None,
                'date_issued': doc.date_issued.isoformat() if doc.date_issued else None,
                'file_name': doc.file_name,
                'file_size': doc.file_size,
                'metadata': doc.metadata,
                'created_at': doc.created_at.isoformat(),
                'updated_at': doc.updated_at.isoformat(),
                'attachments': [a.to_dict() for a in doc.attachments]
            }
        }), 200
    except Exception as e:
        logger.error(f"Error getting document: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@api_bp.route('/upload', methods=['POST'])
def upload_document():
    """Tải lên văn bản"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'}), 400

        # Kiểm tra file extension
        allowed = {'pdf', 'doc', 'docx', 'txt'}
        ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
        if ext not in allowed:
            return jsonify({'success': False, 'error': 'File type not allowed'}), 400

        # Lưu file
        filename = secure_filename(file.filename)
        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S_')
        filename_saved = timestamp + filename
        file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename_saved)
        file.save(file_path)

        # Trích xuất nội dung
        content = extract_file_content(file_path, file.filename)

        # Tạo document
        doc = Document(
            title=request.form.get('title', file.filename),
            content=content,
            document_type=request.form.get('document_type'),
            document_number=request.form.get('document_number'),
            sender=request.form.get('sender'),
            file_path=file_path,
            file_name=file.filename,
            file_size=os.path.getsize(file_path),
            file_type=ext,
            metadata={
                'tags': request.form.get('tags', '').split(',') if request.form.get('tags') else [],
                'priority': request.form.get('priority', 'Normal')
            }
        )

        db.session.add(doc)
        db.session.flush()

        # Xử lý attachments
        attachment_count = 0
        if 'attachments' in request.files:
            for att_file in request.files.getlist('attachments'):
                if att_file and att_file.filename != '':
                    att_filename = secure_filename(att_file.filename)
                    att_timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S_')
                    att_filename_saved = att_timestamp + att_filename
                    att_path = os.path.join(current_app.config['UPLOAD_FOLDER'], att_filename_saved)
                    att_file.save(att_path)

                    attachment = Attachment(
                        document_id=doc.id,
                        filename=att_file.filename,
                        file_path=att_path,
                        file_size=os.path.getsize(att_path),
                        file_type=att_file.filename.rsplit('.', 1)[1].lower() if '.' in att_file.filename else 'unknown'
                    )
                    db.session.add(attachment)
                    attachment_count += 1

        db.session.commit()

        return jsonify({
            'success': True,
            'message': f'Document uploaded successfully',
            'document': doc.to_dict()
        }), 201

    except Exception as e:
        db.session.rollback()
        logger.error(f"Upload error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@api_bp.route('/search', methods=['POST'])
def search_documents():
    """Tìm kiếm văn bản"""
    try:
        data = request.get_json()
        query = data.get('query', '')

        if not query:
            return jsonify({'success': True, 'results': []}), 200

        results = AIService.search_documents(query, limit=10)

        formatted_results = []
        for result in results:
            doc = result['document']
            formatted_results.append({
                'document': {
                    'id': doc.id,
                    'title': doc.title,
                    'document_number': doc.document_number,
                    'document_type': doc.document_type,
                    'sender': doc.sender,
                    'created_at': doc.created_at.isoformat()
                },
                'score': result['score'],
                'matches': result['matches'][:3]
            })

        return jsonify({
            'success': True,
            'results': formatted_results
        }), 200

    except Exception as e:
        logger.error(f"Search error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@api_bp.route('/chat', methods=['POST'])
def chat():
    """Chat với AI"""
    try:
        data = request.get_json()
        message = data.get('message', '')
        session_id = data.get('session_id', 'default')

        if not message:
            return jsonify({'success': False, 'error': 'No message provided'}), 400

        response, related_docs = AIService.process_chat_message(message)

        # Lưu chat message
        chat_msg = ChatMessage(
            session_id=session_id,
            user_message=message,
            ai_response=response,
            related_documents=[d.id for d in related_docs] if related_docs else None
        )
        db.session.add(chat_msg)
        db.session.commit()

        return jsonify({
            'success': True,
            'response': response,
            'results': [
                {
                    'id': d.id,
                    'title': d.title,
                    'snippet': d.content[:150] + '...' if d.content and len(d.content) > 150 else d.content
                } for d in related_docs[:5]
            ]
        }), 200

    except Exception as e:
        logger.error(f"Chat error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@api_bp.route('/download/<doc_id>', methods=['GET'])
def download_document(doc_id):
    """Tải văn bản gốc"""
    try:
        doc = Document.query.get(doc_id)
        if not doc or not doc.file_path:
            return jsonify({'success': False, 'error': 'Document not found'}), 404

        if not os.path.exists(doc.file_path):
            return jsonify({'success': False, 'error': 'File not found'}), 404

        return send_file(
            doc.file_path,
            as_attachment=True,
            download_name=doc.file_name or 'document'
        ), 200

    except Exception as e:
        logger.error(f"Download error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@api_bp.route('/download/attachment/<att_id>', methods=['GET'])
def download_attachment(att_id):
    """Tải file đính kèm"""
    try:
        att = Attachment.query.get(att_id)
        if not att or not att.file_path:
            return jsonify({'success': False, 'error': 'Attachment not found'}), 404

        if not os.path.exists(att.file_path):
            return jsonify({'success': False, 'error': 'File not found'}), 404

        return send_file(
            att.file_path,
            as_attachment=True,
            download_name=att.filename
        ), 200

    except Exception as e:
        logger.error(f"Attachment download error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@api_bp.route('/statistics', methods=['GET'])
def get_statistics():
    """Lấy thống kê hệ thống"""
    try:
        total_docs = Document.query.count()
        total_size = sum([d.file_size or 0 for d in Document.query.all()])

        doc_types = {}
        for doc in Document.query.all():
            dtype = doc.document_type or 'Unknown'
            doc_types[dtype] = doc_types.get(dtype, 0) + 1

        return jsonify({
            'success': True,
            'statistics': {
                'total_documents': total_docs,
                'total_file_size': total_size,
                'total_file_size_mb': total_size / (1024 * 1024),
                'document_types': doc_types
            }
        }), 200

    except Exception as e:
        logger.error(f"Statistics error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


# Export
__all__ = ['api_bp']